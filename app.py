import os
import logging
import qrcode
import subprocess
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

app = Flask(__name__)
app.secret_key = "myytjyujyusupersecretkey"

# Set correct paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOG_PATH = os.path.join(BASE_DIR, "app.log")
STATIC_DIR = os.path.join(BASE_DIR, "static")

# Ensure static directory exists for storing PDFs
os.makedirs(STATIC_DIR, exist_ok=True)

# Configure logging
logging.basicConfig(
    filename=LOG_PATH,
    level=logging.ERROR,
    format="%(asctime)s - %(levelname)s - %(message)s",
)


def convert_to_pdf(input_path):
    """Convert DOCX to PDF using LibreOffice CLI (Works on PythonAnywhere)."""
    output_path = input_path.replace(".docx", ".pdf")
    try:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", input_path, "--outdir", os.path.dirname(input_path)],
            check=True
        )
        return output_path if os.path.exists(output_path) else None
    except subprocess.CalledProcessError as e:
        logging.error(f"Error converting DOCX to PDF: {e}")
        return None


def fill_form(template_path, output_path, data, pdf_url):
    """Fills a Word template with data and converts it to a PDF with a QR code."""
    try:
        doc = Document(template_path)

        # Generate QR Code linking to the PDF
        qr = qrcode.make(pdf_url)
        qr_path = os.path.join(STATIC_DIR, "qr_code.png")
        qr.save(qr_path)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if f"{{{key}}}" in paragraph.text:
                    paragraph.text = paragraph.text.replace(f"{{{key}}}", str(value))

        # Replace placeholders in tables
        qr_inserted = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if f"{{{key}}}" in cell.text:
                            cell.text = cell.text.replace(f"{{{key}}}", str(value))

                    # Insert QR Code if placeholder is found
                    if "{qr_code}" in cell.text:
                        cell.text = ""
                        para = cell.paragraphs[0]
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = para.add_run()
                        run.add_picture(qr_path, width=Inches(1.5))
                        qr_inserted = True

        # If no placeholder for QR code was found, insert at the end
        if not qr_inserted:
            para = doc.add_paragraph("Scan the QR code to view your Gate Pass:")
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = para.add_run()
            run.add_picture(qr_path, width=Inches(2.0))

        # Save Word document
        doc.save(output_path)

        # Convert DOCX to PDF using LibreOffice
        pdf_path = convert_to_pdf(output_path)
        if not pdf_path:
            logging.error("Failed to convert DOCX to PDF.")
            return None

        return pdf_path
    except Exception as e:
        logging.error(f"Error filling form: {e}")
        return None


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/submit", methods=["POST"])
def submit_form():
    """Handles form submission and generates a PDF with a QR code."""
    try:
        # Get form data
        name = request.form.get("name")
        roll_no = request.form.get("roll_no")
        date_from = request.form.get("datefrom")
        date_to = request.form.get("dateto")
        arrival_date = request.form.get("arrivaldate")
        arrival_time = request.form.get("arrivaltime")
        home_address = request.form.get("home_address")
        student_contact_no = request.form.get("student_contact_no")
        parent_name = request.form.get("parent_name")
        parent_contact_no = request.form.get("parent_contact_no")

        if not all([name, roll_no, date_from, date_to, arrival_date, arrival_time]):
            flash("Missing required fields!")
            return redirect(url_for("index"))

        # Create output folder with absolute path
        today = datetime.now().strftime("%Y-%m-%d")
        output_folder = os.path.join(STATIC_DIR, "gatepasses", today)
        os.makedirs(output_folder, exist_ok=True)

        # Define file paths
        template_path = os.path.join(BASE_DIR, "template.docx")
        word_path = os.path.join(output_folder, f"{name}_GatePass.docx")
        pdf_url = f"https://ravinder2115115.pythonanywhere.com/view/{today}/{name}_GatePass.pdf"

        # Fill the Word document and convert it to PDF
        pdf_path = fill_form(template_path, word_path, {
            "name": name,
            "roll_no": roll_no,
            "from": date_from,
            "to": date_to,
            "arrivaldate": arrival_date,
            "arrivaltime": arrival_time,
            "home_address": home_address,
            "student_contact_no": student_contact_no,
            "parent_name": parent_name,
            "parent_contact_no": parent_contact_no,
        }, pdf_url)

        if not pdf_path:
            flash("Error generating the Gate Pass PDF.")
            return redirect(url_for("index"))

        # Redirect to the QR code link
        return redirect(pdf_url)

    except Exception as e:
        logging.error(f"Error processing form: {e}")
        flash("An error occurred.")
        return redirect(url_for("index"))


@app.route("/view/<date>/<filename>")
def view_pdf(date, filename):
    """Displays the generated PDF in the browser."""
    pdf_path = os.path.join(STATIC_DIR, "gatepasses", date, filename)
    if os.path.exists(pdf_path):
        return send_file(pdf_path, mimetype="application/pdf")
    return "File not found", 404


if __name__ == "__main__":
    app.run(debug=True)